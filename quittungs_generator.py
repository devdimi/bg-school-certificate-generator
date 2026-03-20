# -*- coding: utf-8 -*-
"""
Quittungs-Generator für Schulgebühren - Finale, automatisierte Version

**Neue Features:**
- **Automatisches Finden von Input-Dateien:** Das Skript sucht beim Start im
  eigenen Verzeichnis nach `schuelerliste.xlsx`, `preise.xlsx` und
  `Quittung-Template.docx` und füllt die Pfade automatisch aus.
- **Konfigurierbares Schuljahr:** Das Schuljahr wird nicht mehr im Code,
  sondern in der `preise.xlsx` im Tabellenblatt 'Konfiguration' festgelegt.
"""

import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
from docx import Document
import os
import sys # Wird für die Pfad-Ermittlung benötigt
from datetime import datetime
from num2words import num2words

def initialize_paths():
    """
    Sucht nach Standard-Dateinamen und füllt die GUI-Pfade vor.
    Diese Logik funktioniert sowohl für das .py-Skript als auch für die .exe-Datei.
    """
    try:
        if getattr(sys, 'frozen', False):
            # Fall 1: Anwendung läuft als kompilierte .exe (erstellt mit PyInstaller).
            # PyInstaller speichert gebündelte Dateien in einem temporären Ordner,
            # dessen Pfad in `sys._MEIPASS` liegt.
            base_path = os.path.dirname(sys.executable)
        else:
            # Fall 2: Anwendung läuft als normales .py-Skript.
            # Der Basispfad ist das Verzeichnis, in dem das Skript liegt.
            base_path = os.path.dirname(os.path.abspath(__file__))

        # Definiere die Standard-Dateinamen
        default_files = {
            "schuelerliste": "mitglieder_mit_zahlungen_und_zeilen.xlsx",
            "preise": "preise.xlsx",
            "template": "Quittung-Template.docx"
        }

        # Baue die vollen Pfade und überprüfe, ob die Dateien existieren
        path_schueler = os.path.join(base_path, default_files["schuelerliste"])
        if os.path.exists(path_schueler):
            excel_path_var.set(path_schueler)

        path_preise = os.path.join(base_path, default_files["preise"])
        if os.path.exists(path_preise):
            prices_path_var.set(path_preise)
        
        path_template = os.path.join(base_path, default_files["template"])
        if os.path.exists(path_template):
            template_path_var.set(path_template)

        out_dir = os.path.join(base_path, "out")
        output_dir_var.set(out_dir)
            
    except Exception as e:
        print(f"Fehler bei der Initialisierung der Pfade: {e}")


def docx_replace_text(doc, placeholder, replacement):
    def replace_in_paragraph(paragraph):
        # Combine all runs' text
        full_text = ''.join(run.text for run in paragraph.runs)
        if placeholder not in full_text:
            return

        # Replace the placeholder in full text
        new_text = full_text.replace(placeholder, replacement)

        # Clear all runs
        for run in paragraph.runs:
            run.text = ''
        # Set the first run to the replaced text
        if paragraph.runs:
            paragraph.runs[0].text = new_text

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)



def load_prices(filepath):
    """
    Lädt die Preisinformationen UND das Schuljahr aus der Excel-Datei.
    :return: Ein Tupel (Dictionary Kindergebühren, Float Mitgliedsbeitrag, String Schuljahr).
    """
    price_sheets = pd.read_excel(filepath, sheet_name=None)
    
    fee_df = price_sheets['Gebuehren']
    fee_df.columns = fee_df.columns.str.strip() 
    child_fees = pd.Series(fee_df.Betrag.values, index=fee_df.Kind_Nr).to_dict()

    contribution_df = price_sheets['Beitraege']
    contribution_df.columns = contribution_df.columns.str.strip()
    membership_fee = contribution_df[contribution_df.Posten == 'Mitgliedsbeitrag']['Betrag'].iloc[0]
    
    # Neu: Lese das Konfigurationsblatt für das Schuljahr
    config_df = price_sheets['Konfiguration']
    config_df.columns = config_df.columns.str.strip()
    school_year = config_df[config_df.Eigenschaft == 'Schuljahr']['Wert'].iloc[0]
    
    return child_fees, float(membership_fee), str(school_year)


def generate_receipts():
    """Die Hauptfunktion, die den gesamten Generierungsprozess steuert."""
    excel_path = excel_path_var.get()
    template_path = template_path_var.get()
    prices_path = prices_path_var.get()
    output_dir = output_dir_var.get()
    
    if not all([excel_path, template_path, prices_path, output_dir]):
        messagebox.showerror("Fehler", "Bitte alle Pfade auswählen!")
        return

    try:
        # Lade jetzt drei Werte, inklusive des Schuljahres
        child_fees, membership_fee, school_year = load_prices(prices_path)
        

        df = pd.read_excel(excel_path)
        
        df['Mitglied'] = df['Mitglied'].astype(str).str.strip()
        
        # # --- DEBUGGING STEP 2: Inspect unique 'Mitglied' values AFTER cleaning ---
        # print("\n--- Unique 'Mitglied' values AFTER cleaning ---")
        # print(df['Mitglied'].value_counts())
        # print("-" * 40)


        grouped = df.groupby('Mitglied')

        # Get the count of entries for each unique 'Mitglied'
        member_counts = df['Mitglied'].value_counts()

        # Filter for members with more than one entry (i.e., more than one child)
        parents_with_multiple_children = member_counts[member_counts > 1]

        # print("Parents with more than one child:")
        # if not parents_with_multiple_children.empty:
        #     for parent_name, count in parents_with_multiple_children.items():
        #         # You might want to capitalize the name for display if you lowercased it earlier
        #         print(f"- {parent_name.title()} (Number of children: {count})")
        # else:
        #     print("No parents found with more than one child.")

        quittungs_nr = 229

        for (parent_full_name), group in grouped:
            klassen = group['Klasse'].tolist()
            if group['Klasse'].isin(['Abgemeldet', 'Warteliste']).any():
                continue  # Skip this group
            doc = Document(template_path)
            num_children = len(group)
            children_names = " und ".join(group['Kind'])
           
            total_school_fee = sum(child_fees.get(i, 0) for i in range(1, num_children + 1))
            total_amount = total_school_fee + membership_fee
        
            #gezahlt_series = group['Gezahlt']
            #gezahlt_numeric = pd.to_numeric(gezahlt_series, errors='coerce').fillna(0)
            #total_amount_for_member = gezahlt_numeric.sum()
            #if total_amount_for_member > 0:
            #    total_amount = total_amount_for_member
            #    total_school_fee = total_amount - membership_fee

            
            gebuehr_wort = num2words(int(total_school_fee), lang='de')
            mitglied_wort = num2words(int(membership_fee), lang='de')
            gesamt_wort = num2words(int(total_amount), lang='de')

            replacements = {
                "{{ELTERN_NAME}}": parent_full_name,
                "{{KINDER_NAMEN}}": children_names,
                "{{NR}}": f"{quittungs_nr:03d}",
                "{{DATUM}}": datetime.now().strftime("%d.%m.%Y"),
                "{{SCHULJAHR}}": school_year,
                "{{BETRAG_GEBUEHR}}": f"{total_school_fee:,.2f} EUR".replace(",", "X").replace(".", ",").replace("X", "."),
                "{{GESAMTBETRAG}}": f"{total_amount:,.2f} EUR".replace(",", "X").replace(".", ",").replace("X", "."),
                "{{BETRAG_GEBUEHR_WORT}}": f"{gebuehr_wort} Euro",
                "{{GESAMTBETRAG_WORT}}": f"{gesamt_wort} Euro",
                "{{BETRAG_MITGLIED}}": f"{membership_fee:,.2f} EUR".replace(",", "X").replace(".", ",").replace("X", "."),
                "{{BETRAG_MITGLIED_WORT}}": f"{mitglied_wort} Euro",
            }

            for old, new in replacements.items():
                docx_replace_text(doc, old, str(new))

            # ✅ Move this block into the loop
            parent_name = parent_full_name.replace(" ", "_")
            outdir_class = os.path.join(output_dir, klassen[0])
            if not os.path.exists(outdir_class):
                os.makedirs(outdir_class)

            output_filename = os.path.join(outdir_class, f"Quittung_{parent_name}_{quittungs_nr:03d}.docx")
            if os.path.exists(output_filename):
                os.remove(output_filename)
            doc.save(output_filename)
            quittungs_nr += 1

            output = parent_full_name + " " + children_names + " " + klassen[0]
            if len(klassen) > 1 and len(klassen[1].strip()) > 0:
                output += +" " + klassen[1]
            print(output)

            quittungs_nr += 1
        messagebox.showinfo("Erfolg", f"{quittungs_nr - 1} Quittung(en) erfolgreich erstellt!")
    except Exception as e:
        messagebox.showerror("Fehler", f"Ein Fehler ist aufgetreten:\n{e}")

# --- GUI Code ---
def select_excel_file():
    filepath = filedialog.askopenfilename(filetypes=[("Excel-Dateien", "*.xlsx *.xls")])
    if filepath: excel_path_var.set(filepath)

def select_template_file():
    filepath = filedialog.askopenfilename(filetypes=[("Word-Dokumente", "*.docx")])
    if filepath: template_path_var.set(filepath)

def select_prices_file():
    filepath = filedialog.askopenfilename(filetypes=[("Excel-Dateien", "*.xlsx *.xls")])
    if filepath: prices_path_var.set(filepath)

def select_output_dir():
    dirpath = filedialog.askdirectory()
    if dirpath: output_dir_var.set(dirpath)

# Erstelle das Hauptfenster
root = tk.Tk()
root.title("Quittungs-Generator (Auto-Detect Version)")
root.geometry("600x320")

# Erstelle die String-Variablen für die Pfade
excel_path_var, template_path_var, prices_path_var, output_dir_var = tk.StringVar(), tk.StringVar(), tk.StringVar(), tk.StringVar()

# Erstelle den Haupt-Frame
frame = tk.Frame(root, padx=10, pady=10)
frame.pack(expand=True, fill=tk.BOTH)

# Erstelle die GUI-Elemente (Widgets)
tk.Label(frame, text="1. Excel-Datei (Schülerliste) auswählen:").grid(row=0, column=0, sticky="w", pady=2)
tk.Entry(frame, textvariable=excel_path_var, width=60).grid(row=1, column=0, padx=(0, 5))
tk.Button(frame, text="Durchsuchen...", command=select_excel_file).grid(row=1, column=1)
tk.Label(frame, text="2. Excel-Datei (Preise) auswählen:").grid(row=2, column=0, sticky="w", pady=(10, 2))
tk.Entry(frame, textvariable=prices_path_var, width=60).grid(row=3, column=0, padx=(0, 5))
tk.Button(frame, text="Durchsuchen...", command=select_prices_file).grid(row=3, column=1)
tk.Label(frame, text="3. Word-Vorlagendatei auswählen:").grid(row=4, column=0, sticky="w", pady=(10, 2))
tk.Entry(frame, textvariable=template_path_var, width=60).grid(row=5, column=0, padx=(0, 5))
tk.Button(frame, text="Durchsuchen...", command=select_template_file).grid(row=5, column=1)
tk.Label(frame, text="4. Ausgabeordner auswählen:").grid(row=6, column=0, sticky="w", pady=(10, 2))
tk.Entry(frame, textvariable=output_dir_var, width=60).grid(row=7, column=0, padx=(0, 5))
tk.Button(frame, text="Durchsuchen...", command=select_output_dir).grid(row=7, column=1)
tk.Button(frame, text="🚀 Quittungen generieren", font=("Helvetica", 12, "bold"), command=generate_receipts, bg="#4CAF50", fg="white").grid(row=8, column=0, columnspan=2, pady=20, ipadx=10, ipady=5)

# **NEU**: Rufe die Initialisierungsfunktion nach dem Erstellen der GUI auf
initialize_paths()

# Starte die Anwendung
root.mainloop()